from docx import Document
from pathlib import Path

def create_test_docx(output_path):
    """Create a test docx file with sections."""
    doc = Document()
    
    # Add title
    doc.add_heading('Test Document', 0)
    
    # Add sections
    sections = [
        ('Introduction', 'This is the introduction section.'),
        ('Methods', 'This section describes the methods used.'),
        ('Results', 'This section presents the results.'),
        ('Discussion', 'This section discusses the findings.'),
        ('Conclusion', 'This is the conclusion section.')
    ]
    
    for title, content in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
    
    # Save the document
    doc.save(output_path)

if __name__ == '__main__':
    # Create test directory if it doesn't exist
    test_dir = Path('tests/test_files')
    test_dir.mkdir(parents=True, exist_ok=True)
    
    # Create test docx file
    output_path = test_dir / 'sample.docx'
    create_test_docx(str(output_path))
    print(f"Created test docx file at: {output_path}") 
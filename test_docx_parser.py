"""
Tests for the docx_parser module.
"""

import os
import pytest
from pathlib import Path
from docx import Document
from docx_parser import DocxParser, Section

@pytest.fixture
def sample_docx():
    """Create a sample Word document for testing."""
    doc = Document()
    
    # Add a title
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('This is the introduction section.')
    
    # Add a section
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('This is the first paragraph of section 1.')
    doc.add_paragraph('This is the second paragraph of section 1.')
    
    # Add a subsection
    doc.add_heading('Subsection 1.1', level=2)
    doc.add_paragraph('This is a subsection paragraph.')
    
    # Add another section
    doc.add_heading('Section 2', level=1)
    doc.add_paragraph('This is section 2 content.')
    
    # Save the document
    doc_path = Path('test_document.docx')
    doc.save(str(doc_path))
    
    yield doc_path
    
    # Cleanup
    if doc_path.exists():
        doc_path.unlink()

def test_parser_initialization(sample_docx):
    """Test that the parser initializes correctly."""
    parser = DocxParser(str(sample_docx))
    assert parser.file_path == sample_docx
    assert len(parser.elements) > 0

def test_parse_sections(sample_docx):
    """Test that sections are parsed correctly."""
    parser = DocxParser(str(sample_docx))
    sections = parser.parse_sections()
    
    # Should have at least 3 sections (Introduction, Section 1, Section 2)
    assert len(sections) >= 3
    
    # Check first section
    assert sections[0].title == 'Introduction'
    assert len(sections[0].content) > 0
    
    # Check second section
    assert sections[1].title == 'Section 1'
    assert len(sections[1].content) >= 2

def test_get_section_by_title(sample_docx):
    """Test retrieving a section by title."""
    parser = DocxParser(str(sample_docx))
    parser.parse_sections()
    
    section = parser.get_section_by_title('Section 1')
    assert section is not None
    assert section.title == 'Section 1'
    assert len(section.content) >= 2

def test_save_sections_to_files(sample_docx, tmp_path):
    """Test saving sections to files."""
    parser = DocxParser(str(sample_docx))
    parser.parse_sections()
    
    output_dir = tmp_path / 'sections'
    parser.save_sections_to_files(str(output_dir))
    
    # Check that files were created
    assert output_dir.exists()
    files = list(output_dir.glob('*.txt'))
    assert len(files) >= 3  # Should have at least 3 section files

def test_invalid_file():
    """Test handling of invalid file."""
    with pytest.raises(FileNotFoundError):
        DocxParser('nonexistent.docx') 